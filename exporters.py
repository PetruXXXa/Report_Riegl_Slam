"""
Модуль для экспорта отчета в PDF и Word (шаг 9).
"""

import os
import tempfile
import subprocess
import re
import base64
import uuid
from io import BytesIO
from typing import Optional
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Mm, Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_cell_border(cell, border_size=4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    old = tcPr.find(qn('w:tcBorders'))
    if old is not None:
        tcPr.remove(old)
    borders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'bottom', 'right']:
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(border_size))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    tcPr.append(borders)


def _set_cell_margins(cell, top=0, bottom=0, left=30, right=30):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for edge, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)


def _set_row_height(row, height_pt=8):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_pt * 20)))
    trHeight.set(qn('w:hRule'), 'exact')
    trPr.append(trHeight)


def _img_from_base64(src):
    if src.startswith('data:image/png;base64,'):
        return base64.b64decode(src.replace('data:image/png;base64,', ''))
    if src.startswith('data:image/jpeg;base64,') or src.startswith('data:image/jpg;base64,'):
        return base64.b64decode(re.sub(r'^data:image/jpe?g;base64,', '', src))
    return None


class ReportExporter:
    """Экспортирует отчет в различные форматы"""

    def __init__(self):
        self._playwright = None
        self._browser = None

    def _get_browser(self):
        if self._browser is None:
            from playwright.sync_api import sync_playwright
            self._playwright = sync_playwright().start()
            self._browser = self._playwright.chromium.launch()
        return self._browser

    def _close_browser(self):
        if self._browser:
            self._browser.close()
            self._browser = None
        if self._playwright:
            self._playwright.stop()
            self._playwright = None

    def export_to_pdf(self, html_content: str, output_path: str) -> bool:
        """Экспортирует HTML в PDF через DOCX (python-docx → docx2pdf)"""
        import traceback
        try:
            temp_docx = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
            temp_docx_path = temp_docx.name
            temp_docx.close()

            self._html_to_docx(html_content, temp_docx_path)

            from docx2pdf import convert
            convert(temp_docx_path, output_path)

            os.unlink(temp_docx_path)
            return True
        except Exception as e:
            traceback.print_exc()
            print(f"Ошибка экспорта в PDF: {e}")
            return False

    def export_to_jpeg(self, html_content: str, output_base: str) -> bool:
        """Экспортирует HTML в JPEG через DOCX → PDF → PyMuPDF"""
        import traceback
        try:
            temp_docx = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
            temp_docx_path = temp_docx.name
            temp_docx.close()

            temp_pdf = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False)
            temp_pdf_path = temp_pdf.name
            temp_pdf.close()

            self._html_to_docx(html_content, temp_docx_path)

            from docx2pdf import convert
            convert(temp_docx_path, temp_pdf_path)

            import fitz
            doc = fitz.open(temp_pdf_path)
            base, ext = os.path.splitext(output_base)
            for i, page in enumerate(doc, 1):
                pix = page.get_pixmap(dpi=200)
                out_file = f"{base}{ext}_{i:03d}.jpg"
                pix.save(out_file, jpg_quality=75)
                print(f"Сохранено: {out_file}")

            doc.close()
            os.unlink(temp_docx_path)
            os.unlink(temp_pdf_path)
            return True

        except Exception as e:
            traceback.print_exc()
            print(f"Ошибка экспорта в JPEG: {e}")
            return False

    def export_to_word(self, html_content: str, output_path: str) -> bool:
        """Экспортирует HTML в DOCX через python-docx с полным контролем форматирования"""
        import traceback
        try:
            self._html_to_docx(html_content, output_path)
            return True
        except Exception as e:
            traceback.print_exc()
            print(f"Ошибка экспорта в Word: {e}")
            return False

    def export_to_word_enhanced(self, html_content: str, output_path: str) -> bool:
        """Экспортирует HTML в Word"""
        return self.export_to_word(html_content, output_path)

    def _html_to_docx(self, html_content: str, output_path: str):
        doc = Document()

        section = doc.sections[0]
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)
        section.left_margin = Mm(20)
        section.right_margin = Mm(15)

        soup = BeautifulSoup(html_content, 'html.parser')
        import traceback

        title = soup.find('h1')
        if title:
            h = doc.add_heading(title.get_text(strip=True), level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        def _add_table(table_tag):
            rows = table_tag.find_all('tr')
            if not rows:
                return
            cols = max(len(r.find_all(['th', 'td'])) for r in rows)
            if cols == 0:
                return
            doc_table = doc.add_table(rows=len(rows), cols=cols)
            doc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            doc_table.autofit = True
            for i, row in enumerate(rows):
                cells = row.find_all(['th', 'td'])
                doc_row = doc_table.rows[i]
                _set_row_height(doc_row, 14)
                for j, cell in enumerate(cells):
                    if j >= cols:
                        break
                    doc_cell = doc_row.cells[j]
                    doc_cell.text = ''
                    text = cell.get_text(strip=True) or ''
                    if text:
                        run = doc_cell.paragraphs[0].add_run(text)
                        run.font.size = Pt(9)
                        if cell.name == 'th':
                            run.bold = True
                    _set_cell_border(doc_cell, border_size=4)
                    _set_cell_margins(doc_cell, top=0, bottom=0, left=30, right=30)
                    doc_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        def _add_img(img_tag, width_mm=165):
            src = img_tag.get('src', '')
            data = _img_from_base64(src)
            if data:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                try:
                    p.add_run().add_picture(BytesIO(data), width=Mm(width_mm))
                except Exception:
                    pass

        def _process_element(el):
            if el.name == 'h2':
                doc.add_heading(el.get_text(strip=True), level=2)
                doc.paragraphs[-1].paragraph_format.space_before = Pt(6)
                doc.paragraphs[-1].paragraph_format.space_after = Pt(4)
            elif el.name == 'h3':
                doc.add_heading(el.get_text(strip=True), level=3)
                doc.paragraphs[-1].paragraph_format.space_before = Pt(4)
                doc.paragraphs[-1].paragraph_format.space_after = Pt(2)
            elif el.name == 'h4':
                doc.add_heading(el.get_text(strip=True), level=4)
                doc.paragraphs[-1].paragraph_format.space_before = Pt(3)
                doc.paragraphs[-1].paragraph_format.space_after = Pt(1)
            elif el.name == 'table' and 'accuracy-table' in el.get('class', []):
                _add_table(el)
            elif el.name == 'img':
                _add_img(el, 165)
            elif el.name == 'div':
                cls = el.get('class', [])
                if 'subsection' in cls:
                    for child in el.children:
                        if child.name:
                            _process_element(child)
                elif 'text-stats' in cls:
                    for item in el.find_all('div', class_='text-stat-item'):
                        label = item.find('span', class_='text-stat-label')
                        value = item.find('span', class_='text-stat-value')
                        if label:
                            p = doc.add_paragraph()
                            p.paragraph_format.space_before = Pt(0)
                            p.paragraph_format.space_after = Pt(0)
                            p.paragraph_format.line_spacing = Pt(13)
                            run_l = p.add_run(label.get_text(strip=True))
                            run_l.bold = True
                            run_l.font.size = Pt(11)
                            if value:
                                run_v = p.add_run('  ' + value.get_text(strip=True))
                                run_v.font.size = Pt(11)
                elif 'map-container' in cls:
                    img_tag = el.find('img')
                    if img_tag:
                        _add_img(img_tag, 170)
                elif 'image-container' in cls:
                    img_tag = el.find('img')
                    if img_tag:
                        _add_img(img_tag, 165)
                elif 'table-container' in cls:
                    for child in el.children:
                        if child.name:
                            _process_element(child)
            elif el.name == 'p':
                p = doc.add_paragraph(el.get_text(strip=True))
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                for run in p.runs:
                    run.font.size = Pt(10)

        try:
            sections = soup.find_all('div', class_='section')
            for idx, section_div in enumerate(sections):
                cls = section_div.get('class', [])
                if idx > 0 and 'page-break-before' in cls:
                    doc.add_page_break()
                for child in section_div.children:
                    if child.name:
                        _process_element(child)
        except Exception:
            traceback.print_exc()
            raise

        doc.save(output_path)

    def _is_pandoc_available(self) -> bool:
        """Проверяет, доступен ли pandoc в системе"""
        try:
            result = subprocess.run(['pandoc', '--version'], capture_output=True, text=True)
            return result.returncode == 0
        except (FileNotFoundError, OSError):
            return False

    def _export_to_mhtml(self, html_content: str, output_path: str) -> bool:
        """Экспортирует HTML в MHTML формат (MIME HTML)"""
        try:
            html_modified = html_content
            images = []
            pattern = r'<img[^>]+src="data:image/([^;]+);base64,([^"]+)"'

            def replace_match(match):
                img_type = match.group(1)
                b64_data = match.group(2)
                content_type_map = {
                    'png': 'image/png', 'jpeg': 'image/jpeg', 'jpg': 'image/jpeg',
                    'gif': 'image/gif', 'svg+xml': 'image/svg+xml'
                }
                content_type = content_type_map.get(img_type, f'image/{img_type}')
                cid = f"image{len(images)+1}@report"
                images.append((content_type, b64_data, cid))
                return f'<img src="{cid}" alt="image"'

            html_modified = re.sub(pattern, replace_match, html_modified, flags=re.IGNORECASE)
            boundary = f"----=_NextPart_{uuid.uuid4().hex[:16]}"

            with open(output_path, 'w', encoding='utf-8') as f:
                f.write("From: <Saved by Report Generator>\n")
                f.write("Subject: Report\n")
                f.write("MIME-Version: 1.0\n")
                f.write(f"Content-Type: multipart/related; type=\"text/html\"; boundary=\"{boundary}\"\n\n")
                f.write(f"--{boundary}\n")
                f.write("Content-Type: text/html; charset=utf-8\n")
                f.write("Content-Transfer-Encoding: 8bit\n")
                f.write("Content-Location: file:///report.html\n\n")
                f.write(html_modified + "\n")
                for content_type, data_b64, cid in images:
                    f.write(f"--{boundary}\n")
                    f.write(f"Content-Type: {content_type}\n")
                    f.write("Content-Transfer-Encoding: base64\n")
                    f.write(f"Content-Location: {cid}\n\n")
                    for j in range(0, len(data_b64), 76):
                        f.write(data_b64[j:j+76] + "\n")
                    f.write("\n")
                f.write(f"--{boundary}--\n")
            return True
        except Exception as e:
            print(f"Ошибка экспорта в MHTML: {e}")
            return False

    def export_with_pandoc(self, html_content: str, output_path: str, format: str = 'docx') -> bool:
        """Экспорт с использованием pandoc (более качественный)"""
        try:
            temp_html = tempfile.NamedTemporaryFile(suffix='.html', delete=False, mode='w', encoding='utf-8')
            temp_html_path = temp_html.name
            temp_html.write(html_content)
            temp_html.close()

            cmd = ['pandoc', temp_html_path, '-o', output_path]
            if format == 'docx':
                cmd.extend(['--to', 'docx'])
                cmd.extend(['-V', 'geometry:top=20mm'])
                cmd.extend(['-V', 'geometry:bottom=20mm'])
                cmd.extend(['-V', 'geometry:left=20mm'])
                cmd.extend(['-V', 'geometry:right=15mm'])
                cmd.extend(['-V', 'geometry:paper=a4'])
            elif format == 'pdf':
                cmd.extend(['--to', 'pdf', '--pdf-engine=wkhtmltopdf'])

            result = subprocess.run(cmd, capture_output=True, text=True)
            os.unlink(temp_html_path)
            return result.returncode == 0
        except Exception as e:
            print(f"Ошибка экспорта с pandoc: {e}")
            return False
